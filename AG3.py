# utf-8
import os
os.environ['PYTHONIOENCODING'] = 'utf-8'
os.environ['PYTHONUTF8'] = '1'

import os
import numpy as np
import streamlit as st
from openai import OpenAI
import chardet  # 新增：自动检测编码

# ------------------------- LangChain 基础导入（架构用）-------------------------
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.prompts import PromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_openai import ChatOpenAI
from langchain_core.embeddings import Embeddings 
from langchain_core.output_parsers import StrOutputParser

# ------------------------- 【手写核心1】自定义通义千问 Embedding-------------------------
class QwenCustomEmbedding(Embeddings):
    def __init__(self, api_key):
        super().__init__()
        self.client = OpenAI(
            api_key=api_key,
            base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"
        )

    def embed_documents(self, texts):
        embeddings = []
        for text in texts:
            resp = self.client.embeddings.create(
                model="text-embedding-v1",
                input=text
            )
            embeddings.append(resp.data[0].embedding)
        return embeddings

    def embed_query(self, text):
        resp = self.client.embeddings.create(
            model="text-embedding-v1",
            input=text
        )
        return resp.data[0].embedding

# ------------------------- 【手写核心2】增强的文本切分-------------------------
def custom_text_splitter(text: str, chunk_size=300, overlap=50):
    # 🔧 修复乱码：清理文本
    text = text.encode('utf-8', errors='ignore').decode('utf-8')
    
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", "。", "！", "？", " ", "."]
    )
    chunks = splitter.split_text(text)
    return [Document(page_content=chunk) for chunk in chunks]

# ------------------------- 【手写核心3】自定义余弦相似度检索-------------------------
def cosine_similarity(vec1, vec2):
    dot = np.dot(vec1, vec2)
    norm1 = np.linalg.norm(vec1)
    norm2 = np.linalg.norm(vec2)
    if norm1 == 0 or norm2 == 0:
        return 0.0
    return dot / (norm1 * norm2)

# ------------------------- 【手写核心4】增强的强约束 Prompt-------------------------
CUSTOM_PROMPT = PromptTemplate(
    input_variables=["context", "question"],
    template="""你是企业知识库智能助手，必须严格遵守以下规则：
1. 只使用提供的上下文信息回答，绝对不能编造
2. 回答要完整、通顺、正式，使用中文
3. 上下文无答案时，直接回复：无法从知识库找到答案
4. 如果上下文包含乱码，尝试理解并给出最佳答案
5. 不要解释、不要多余内容

上下文：
{context}

问题：{question}
回答："""
)

# ------------------------- 【手写核心5】自定义大模型调用-------------------------
class QwenChatModel:
    def __init__(self, api_key):
        self.client = OpenAI(
            api_key=api_key,
            base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"
        )

    def invoke(self, prompt):
        completion = self.client.chat.completions.create(
            model="qwen-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0
        )
        return completion.choices[0].message.content.strip()

# ------------------------- 增强的文件读取函数-------------------------
def read_uploaded_file(uploaded_file):
    """安全读取上传文件，自动处理编码问题"""
    try:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension == 'txt':
            # 🔧 自动检测编码
            raw_data = uploaded_file.read()
            result = chardet.detect(raw_data)
            encoding = result['encoding'] if result['encoding'] else 'utf-8'
            return raw_data.decode(encoding, errors='ignore')
            
        elif file_extension == 'pdf':
            import PyPDF2
            import io
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
            content = ""
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    # 🔧 清理PDF提取的文本
                    text = text.encode('utf-8', errors='ignore').decode('utf-8')
                    content += text + "\n"
            return content
            
        elif file_extension == 'docx':
            import docx
            import io
            doc = docx.Document(io.BytesIO(uploaded_file.read()))
            content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    content += paragraph.text + "\n"
            return content
    
    except Exception as e:
        st.error(f"❌ 文件读取失败: {str(e)}")
        return ""

# ======================== 前端界面 ========================
st.set_page_config(
    page_title="RAG求职演示项目", 
    page_icon="🏢",
    layout="wide"
)

# 自定义CSS样式
st.markdown("""
<style>
    .main-title {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f77b4;
    }
    .success-box {
        padding: 20px;
        border-radius: 10px;
        background-color: #f0f8ff;
        border-left: 5px solid #4CAF50;
    }
</style>
""", unsafe_allow_html=True)

st.markdown('<p class="main-title">🏢 企业级 RAG 知识库问答系统</p>', unsafe_allow_html=True)
st.caption("🚀 技术栈：LangChain + 通义千问 + 手写核心模块 | 求职面试演示项目")

# 添加使用说明
with st.sidebar.expander("📖 使用说明", expanded=True):
    st.markdown("""
    ### 🚀 快速开始
    1. **🔑 输入API密钥**: DashScope API密钥
    2. **📚 准备知识库**: 
       - 📁 上传文件（TXT/PDF/DOCX）
       - 📝 或粘贴文本内容
    3. **❓ 提问**: 输入你的问题
    4. **✨ 获取答案**: 点击"开始回答"
    
    ### 📋 支持格式
    - **TXT**: 纯文本（自动检测编码）
    - **PDF**: PDF文档
    - **DOCX**: Word文档
    
    ### 💡 小贴士
    - 文本质量越高，答案越准确
    - 支持中文、英文混合内容
    - 每次可处理最多10000字
    """)

# 输入API Key
api_key = st.text_input("🔑 请输入 DASHSCOPE_API_KEY", type="password", 
                       help="从 https://dashscope.aliyun.com/ 获取")

# 知识库输入区域
tab1, tab2 = st.tabs(["📁 上传文件", "📝 粘贴文本"])

with tab1:
    uploaded_file = st.file_uploader(
        "选择知识库文件", 
        type=['txt', 'pdf', 'docx'],
        help="支持TXT、PDF、DOCX格式，文件大小不超过10MB"
    )
    
    if uploaded_file is not None:
        with st.spinner("📖 正在读取文件..."):
            content = read_uploaded_file(uploaded_file)
            if content:
                st.session_state.knowledge_content = content
                st.success(f"✅ 已加载: {uploaded_file.name} ({len(content)} 字符)")
                
                # 预览内容
                with st.expander("👀 预览内容"):
                    st.text(content[:500] + "..." if len(content) > 500 else content)

with tab2:
    knowledge_input = st.text_area(
        "在此粘贴知识库文本",
        height=200,
        placeholder="请粘贴你的知识库内容...\n\n例如：\n通义千问是阿里云推出的AI模型...\nAPI密钥格式为：sk-..."
    )
    if knowledge_input:
        st.session_state.knowledge_content = knowledge_input

# 使用session state中的内容
if 'knowledge_content' in st.session_state:
    knowledge_input = st.session_state.knowledge_content

user_question = st.text_input("❓ 请输入你的问题", placeholder="例如：通义千问的API是什么？")

# 运行按钮
col1, col2, col3 = st.columns([1, 2, 1])
with col2:
    run_button = st.button("🚀 开始回答", use_container_width=True)

if run_button:
    if not api_key:
        st.error("❌ 请输入API密钥")
    elif not knowledge_input:
        st.error("❌ 请提供知识库内容")
    elif not user_question:
        st.error("❌ 请输入问题")
    else:
        with st.spinner("🤔 AI正在思考..."):
            try:
                # 1. 切分文档
                docs = custom_text_splitter(knowledge_input)
                
                # 显示处理进度
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                status_text.text(f"📄 文档已切分为 {len(docs)} 块")
                progress_bar.progress(30)
                
                # 2. Embedding
                embed = QwenCustomEmbedding(api_key)
                progress_bar.progress(50)
                
                # 3. 向量库
                db = FAISS.from_documents(docs, embed)
                retriever = db.as_retriever(search_kwargs={"k": 3})
                progress_bar.progress(70)
                
                # 4. 构建RAG链
                llm = ChatOpenAI(
                    api_key=api_key,
                    base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
                    model="qwen-turbo",
                    temperature=0.0
                )

                def format_docs(docs):
                    return "\n".join([d.page_content for d in docs])

                rag_chain = (
                    {
                        "context": retriever | format_docs,
                        "question": RunnablePassthrough()
                    }
                    | CUSTOM_PROMPT
                    | llm
                    | StrOutputParser()
                )
                progress_bar.progress(90)
                
                # 5. 获取答案
                answer = rag_chain.invoke(user_question)
                progress_bar.progress(100)
                status_text.empty()
                
                # 展示结果
                st.subheader("✅ 回答结果")
                st.markdown(f'<div class="success-box">{answer}</div>', unsafe_allow_html=True)
                
                # 显示检索片段
                with st.expander("🔍 查看检索到的相关文档片段"):
                    retrieved_docs = retriever.invoke(user_question)
                    for i, doc in enumerate(retrieved_docs, 1):
                        st.markdown(f"**📄 片段 {i}:**")
                        st.info(doc.page_content)
                        
                        # 计算相似度
                        if i > 1:
                            doc_vec = embed.embed_query(doc.page_content)
                            query_vec = embed.embed_query(user_question)
                            sim_score = cosine_similarity(query_vec, doc_vec)
                            st.caption(f"相似度: {sim_score:.3f}")
                
            except Exception as e:
                st.error(f"❌ 处理出错: {str(e)}")
                st.info("💡 建议检查API密钥是否正确，或尝试简化知识库内容")